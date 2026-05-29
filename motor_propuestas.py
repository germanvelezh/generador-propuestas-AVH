# -*- coding: utf-8 -*-
"""
Motor de propuestas comerciales (Fase B).
- compute_budget(cfg): calcula el presupuesto completo desde la configuracion.
- build_excel(cfg, path): genera el Excel parametrizado (formulas vivas) seedeado con la config.
- build_docx(cfg, budget, path): genera el documento de propuesta para el cliente (Word).
Diseno: una sola fuente (la config) alimenta Excel y documento, garantizando coherencia.
"""
import json, copy
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.worksheet.datavalidation import DataValidation
from openpyxl.utils import get_column_letter

# ============================ CONFIG POR DEFECTO ============================
DEFAULT_CONFIG = {
  "proyecto": {
    "codigo": "CF-XXXXXXXX-XX-01-OC",
    "objeto": "OFERTA TECNICA Y ECONOMICA PARA CARACTERIZACION DE FLORA",
    "cliente": "",
    "fecha": "2026-05-29",
    "n_titulos": 5
  },
  "comercial": {"utilidad": 0.25, "iva": 0.19},
  "desembolsos": [
    {"n": 1, "pct": 0.30, "obs": "Para inicio de actividades"},
    {"n": 2, "pct": 0.30, "obs": "Una vez finalizada la fase de campo"},
    {"n": 3, "pct": 0.30, "obs": "Contra entrega de documentos finales"},
    {"n": 4, "pct": 0.10, "obs": "Contra entrega de ajustes a satisfaccion"}
  ],
  "catalogo": {
    "cargos": [
      {"nombre": "Coordinador(a) general", "salario": 6831500, "factor": 1.59},
      {"nombre": "Profesional flora",      "salario": 5255000, "factor": 1.30},
      {"nombre": "Profesional fauna",      "salario": 5255000, "factor": 1.59},
      {"nombre": "Profesional SIG",        "salario": 5780500, "factor": 1.59},
      {"nombre": "Profesional Social",     "salario": 5780500, "factor": 1.59},
      {"nombre": "Auxiliar Tecnico",       "salario": 3047900, "factor": 1.30},
      {"nombre": "Auxiliar de campo",      "salario": 1576500, "factor": 1.30}
    ],
    "logistica": [
      {"nombre": "Transporte mayor", "unidad": "Trayecto", "valor": 900000},
      {"nombre": "Transporte menor", "unidad": "Trayecto", "valor": 80000},
      {"nombre": "Hospedaje", "unidad": "Noche", "valor": 70000},
      {"nombre": "Alimentacion", "unidad": "dia/persona", "valor": 70000},
      {"nombre": "Hidratacion y refrigerios", "unidad": "Unidad", "valor": 25000}
    ]
  },
  "actividades": [
    {"nombre": "LEVANTAMIENTO INFORMACION PRIMARIA FLORA", "activa": True,
     "tps_campo": 90, "tps_entregables": 60,
     "personal": [
       {"cargo": "Coordinador(a) general", "dedicacion": 1.29, "cantidad": 1},
       {"cargo": "Profesional flora", "dedicacion": 1.0, "cantidad": 4},
       {"cargo": "Auxiliar Tecnico", "dedicacion": 1.0, "cantidad": 4},
       {"cargo": "Profesional SIG", "dedicacion": 2.57, "cantidad": 1},
       {"cargo": "Auxiliar de campo", "dedicacion": 1.0, "cantidad": 4}
     ],
     "logistica": [
       {"concepto": "Transporte mayor", "cantidad": 6, "valor": "Transporte mayor"},
       {"concepto": "Hidratacion y refrigerios", "cantidad": 1080, "valor": "Hidratacion y refrigerios"},
       {"concepto": "Materiales e insumos (global)", "cantidad": 1, "valor": 2138800},
       {"concepto": "Equipos, software y papeleria (global)", "cantidad": 1, "valor": 9624600},
       {"concepto": "EPP (global)", "cantidad": 1, "valor": 1600000}
     ]},
    {"nombre": "LEVANTAMIENTO INFORMACION PRIMARIA FAUNA / HIDROBIOLOGICOS", "activa": False,
     "tps_campo": 9, "tps_entregables": 30,
     "personal": [
       {"cargo": "Coordinador(a) general", "dedicacion": 0.5, "cantidad": 1},
       {"cargo": "Profesional fauna", "dedicacion": 1.0, "cantidad": 3},
       {"cargo": "Profesional SIG", "dedicacion": 1.3, "cantidad": 1},
       {"cargo": "Auxiliar de campo", "dedicacion": 0.3, "cantidad": 2}
     ],
     "logistica": [
       {"concepto": "Transporte mayor", "cantidad": 2, "valor": "Transporte mayor"},
       {"concepto": "Hidratacion y refrigerios", "cantidad": 27, "valor": "Hidratacion y refrigerios"},
       {"concepto": "Materiales e insumos (global)", "cantidad": 1, "valor": 300000},
       {"concepto": "Equipos, software y papeleria (global)", "cantidad": 1, "valor": 2070000},
       {"concepto": "EPP (global)", "cantidad": 1, "valor": 600000}
     ]},
    {"nombre": "PRIMATES", "activa": False,
     "tps_campo": 15, "tps_entregables": 45,
     "personal": [
       {"cargo": "Coordinador(a) general", "dedicacion": 0.4, "cantidad": 1},
       {"cargo": "Profesional fauna", "dedicacion": 1.0, "cantidad": 2},
       {"cargo": "Profesional SIG", "dedicacion": 1.3, "cantidad": 1},
       {"cargo": "Auxiliar Tecnico", "dedicacion": 0.3, "cantidad": 1}
     ],
     "logistica": [
       {"concepto": "Transporte mayor", "cantidad": 2, "valor": "Transporte mayor"},
       {"concepto": "Hidratacion y refrigerios", "cantidad": 18, "valor": "Hidratacion y refrigerios"},
       {"concepto": "Materiales e insumos (global)", "cantidad": 1, "valor": 100000},
       {"concepto": "Equipos, software y papeleria (global)", "cantidad": 1, "valor": 60000}
     ]}
  ],
  "entregables": {"activa": True, "items": [
     {"nombre": "Coordinacion general de entregables", "cargo": "Coordinador(a) general", "dedicacion": 1.0, "cantidad": 1, "tiempo_meses": 1.0},
     {"nombre": "Clasificacion y caracterizacion de coberturas y ecosistemas", "cargo": "Profesional flora", "dedicacion": 1.0, "cantidad": 1, "tiempo_meses": 0.6},
     {"nombre": "Delimitacion y definicion de area de influencia flora", "cargo": "Profesional flora", "dedicacion": 1.0, "cantidad": 1, "tiempo_meses": 0.6},
     {"nombre": "Caracterizacion de flora arborea", "cargo": "Profesional flora", "dedicacion": 1.0, "cantidad": 1, "tiempo_meses": 3.0},
     {"nombre": "Caracterizacion de flora epifita en veda", "cargo": "Profesional flora", "dedicacion": 1.0, "cantidad": 1, "tiempo_meses": 3.0},
     {"nombre": "Plan de compensacion del componente biotico y fichas de manejo", "cargo": "Profesional flora", "dedicacion": 1.0, "cantidad": 1, "tiempo_meses": 1.5},
     {"nombre": "GDB y Cartografia", "cargo": "Profesional flora", "dedicacion": 1.29, "cantidad": 0.7, "tiempo_meses": 1.0}
  ]}
}

# ============================ MOTOR DE CALCULO ============================
import numbers
def _num(x):
    """Convierte a float de forma tolerante (None/'' -> 0); soporta tipos numpy."""
    if x is None or x == "":
        return 0.0
    if isinstance(x, numbers.Number):
        return float(x)
    try:
        return float(str(x).replace(",", ""))
    except ValueError:
        return 0.0

def _cargo_costo(cfg, nombre):
    for c in cfg["catalogo"]["cargos"]:
        if c["nombre"] == nombre:
            return _num(c["salario"]) * _num(c["factor"])
    raise ValueError(f"Cargo no encontrado en catalogo: {nombre}")

def _log_valor(cfg, ref):
    # numero (incluye numpy) -> valor directo (monto global)
    if isinstance(ref, numbers.Number):
        return float(ref)
    s = str(ref).strip()
    for l in cfg["catalogo"]["logistica"]:
        if l["nombre"] == s:
            return _num(l["valor"])
    # texto numerico (p.ej. "2138800")
    try:
        return float(s.replace(",", ""))
    except ValueError:
        raise ValueError(f"Concepto logistico no encontrado: {ref}")

def compute_budget(cfg):
    u = cfg["comercial"]["utilidad"]; iva = cfg["comercial"]["iva"]
    ntit = cfg["proyecto"]["n_titulos"]
    bloques = []
    # actividades de campo
    for act in cfg["actividades"]:
        tm = _num(act["tps_campo"]) / 30.0
        sub = 0.0
        for p in act["personal"]:
            if not p.get("cargo"): continue
            sub += _num(p["dedicacion"]) * _num(p["cantidad"]) * tm * _cargo_costo(cfg, p["cargo"])
        for lg in act["logistica"]:
            if lg.get("concepto") in (None, ""): continue
            sub += _num(lg["cantidad"]) * _log_valor(cfg, lg["valor"])
        bloques.append({"nombre": act["nombre"], "activa": bool(act["activa"]),
                        "subtotal": sub, "utilidad": sub * u, "subtotal_u": sub * (1 + u)})
    # entregables
    ent = cfg["entregables"]
    sub = 0.0
    for it in ent["items"]:
        if not it.get("cargo"): continue
        sub += _num(it["dedicacion"]) * _num(it["cantidad"]) * _num(it["tiempo_meses"]) * _cargo_costo(cfg, it["cargo"])
    bloques.append({"nombre": "Entregables", "activa": bool(ent.get("activa", True)),
                    "subtotal": sub, "utilidad": sub * u, "subtotal_u": sub * (1 + u)})
    cd = sum(b["subtotal"] for b in bloques if b["activa"])
    cu = sum(b["subtotal_u"] for b in bloques if b["activa"])
    iva_val = cu * iva
    total = cu + iva_val
    desem = []
    for d in cfg["desembolsos"]:
        pct = _num(d["pct"]); antes = cu * pct
        desem.append({"n": d.get("n"), "pct": pct, "obs": d.get("obs", ""),
                      "antes_iva": antes, "iva": antes * iva, "total": antes * (1 + iva)})
    return {"bloques": bloques, "cd": cd, "cu": cu, "iva": iva_val, "total": total,
            "iva_pct": iva, "n_titulos": ntit, "desembolsos": desem}

# ============================ EXCEL ============================
NAVY="1F3864"; BLUE="2E75B6"; LIGHT="D9E2F3"; GREY="F2F2F2"; INPUT="FFF2CC"; GREEN="E2EFDA"
fN=Font(size=11); fB=Font(size=11,bold=True); fH=Font(size=11,bold=True,color="FFFFFF")
fW=Font(size=11,bold=True,color="FFFFFF")
fill=lambda c:PatternFill("solid",fgColor=c)
thin=Side(style="thin",color="BFBFBF"); BORD=Border(thin,thin,thin,thin)
CEN=Alignment(horizontal="center",vertical="center",wrap_text=True)
LEF=Alignment(horizontal="left",vertical="center",wrap_text=True)
RIG=Alignment(horizontal="right",vertical="center")
MONEY='#,##0'; PCT='0.0%'; NUM='#,##0.00'

def _set(ws,coord,val,font=fN,bg=None,fmt=None,al=LEF,bd=True):
    c=ws[coord]; c.value=val; c.font=font
    if bg: c.fill=fill(bg)
    if fmt: c.number_format=fmt
    c.alignment=al
    if bd: c.border=BORD
    return c

def build_excel(cfg, path):
    wb=openpyxl.Workbook()
    # ---- Catalogo ----
    cat=wb.active; cat.title="Catalogo"; cat.sheet_view.showGridLines=False
    _set(cat,"A1","CATALOGO DE COSTOS",Font(size=14,bold=True,color=NAVY),bd=False)
    for i,h in enumerate(["Cargo","Salario base (mes)","Factor prestacional","Costo mes ($)","Costo dia ($)","Costo hora ($)"]):
        _set(cat,f"{get_column_letter(i+1)}3",h,fH,BLUE,al=CEN)
    SAL={}; r=4
    for cg in cfg["catalogo"]["cargos"]:
        _set(cat,f"A{r}",cg["nombre"]); _set(cat,f"B{r}",cg["salario"],fN,INPUT,MONEY,RIG)
        _set(cat,f"C{r}",cg["factor"],fN,INPUT,NUM,CEN)
        _set(cat,f"D{r}",f"=B{r}*C{r}",fB,None,MONEY,RIG)
        _set(cat,f"E{r}",f"=D{r}/30",fN,None,MONEY,RIG); _set(cat,f"F{r}",f"=E{r}/8",fN,None,MONEY,RIG)
        SAL[cg["nombre"]]=r; r+=1
    lr=r+1
    _set(cat,f"A{lr}","Costos logisticos y viaticos (valor unitario)",fB,LIGHT); cat.merge_cells(f"A{lr}:F{lr}")
    lr+=1
    for i,h in enumerate(["Concepto","Unidad","Valor unitario ($)"]): _set(cat,f"{get_column_letter(i+1)}{lr}",h,fH,BLUE,al=CEN)
    LOG={}; lr+=1
    for lg in cfg["catalogo"]["logistica"]:
        _set(cat,f"A{lr}",lg["nombre"]); _set(cat,f"B{lr}",lg["unidad"],fN,al=CEN); _set(cat,f"C{lr}",lg["valor"],fN,INPUT,MONEY,RIG)
        LOG[lg["nombre"]]=lr; lr+=1
    cat.column_dimensions["A"].width=34
    for col in "BCDEF": cat.column_dimensions[col].width=16
    CC=lambda n:f"Catalogo!$D${SAL[n]}"; CL=lambda n:f"Catalogo!$C${LOG[n]}"

    # ---- Parametros ----
    par=wb.create_sheet("Parametros"); par.sheet_view.showGridLines=False
    _set(par,"A1","PARAMETROS DEL PROYECTO",Font(size=14,bold=True,color=NAVY),bd=False)
    pj=cfg["proyecto"]
    rows=[("Codigo del proyecto",pj["codigo"]),("Objeto / Nombre de la oferta",pj["objeto"]),
          ("Cliente",pj["cliente"]),("Fecha",pj["fecha"]),("Numero de titulos mineros",pj["n_titulos"])]
    r=3
    for et,val in rows:
        _set(par,f"A{r}",et,fB); _set(par,f"B{r}",val,fN,INPUT); r+=1
    NTIT_ROW=7  # numero titulos
    _set(par,"A9","Utilidad esperada (%)",fB); _set(par,"B9",cfg["comercial"]["utilidad"],fB,INPUT,PCT,CEN)
    _set(par,"A10","IVA (%)",fB); _set(par,"B10",cfg["comercial"]["iva"],fB,INPUT,PCT,CEN)
    UTIL="Parametros!$B$9"; IVA="Parametros!$B$10"; NTIT=f"Parametros!$B${NTIT_ROW}"
    _set(par,"A12","Actividades y TPS (dias)",fB,LIGHT); par.merge_cells("A12:D12")
    for i,h in enumerate(["Actividad","Activa (Si/No)","Dias Fase de campo","Dias Fase de entregables"]):
        _set(par,f"{get_column_letter(i+1)}13",h,fH,BLUE,al=CEN)
    ACT_ROW={}; r=14
    for act in cfg["actividades"]:
        _set(par,f"A{r}",act["nombre"]); _set(par,f"B{r}","Si" if act["activa"] else "No",fB,INPUT,al=CEN)
        _set(par,f"C{r}",act["tps_campo"],fN,INPUT,al=CEN); _set(par,f"D{r}",act["tps_entregables"],fN,INPUT,al=CEN)
        ACT_ROW[act["nombre"]]=r; r+=1
    dv=DataValidation(type="list",formula1='"Si,No"',allow_blank=False); par.add_data_validation(dv)
    for rr in ACT_ROW.values(): dv.add(par[f"B{rr}"])
    dr=r+1
    _set(par,f"A{dr}","Esquema de desembolsos",fB,LIGHT); par.merge_cells(f"A{dr}:D{dr}"); dr+=1
    for i,h in enumerate(["Desembolso","%","Observacion"]): _set(par,f"{get_column_letter(i+1)}{dr}",h,fH,BLUE,al=CEN)
    par.merge_cells(f"C{dr}:D{dr}"); dr+=1; DES_START=dr
    for d in cfg["desembolsos"]:
        _set(par,f"A{dr}",d["n"],fN,al=CEN); _set(par,f"B{dr}",d["pct"],fB,INPUT,PCT,CEN)
        _set(par,f"C{dr}",d["obs"]); par.merge_cells(f"C{dr}:D{dr}"); dr+=1
    DES_END=dr-1
    par.column_dimensions["A"].width=44
    for col in "BCD": par.column_dimensions[col].width=18
    TOG=lambda n:f'IF(Parametros!$B${ACT_ROW[n]}="Si",1,0)'
    TMC=lambda n:f"Parametros!$C${ACT_ROW[n]}/30"

    # ---- Presupuesto ----
    pre=wb.create_sheet("Presupuesto"); pre.sheet_view.showGridLines=False
    _set(pre,"A1","PRESUPUESTO DE EJECUCION",Font(size=14,bold=True,color=NAVY),bd=False)
    for c,w in {"A":40,"B":13,"C":11,"D":14,"E":16,"F":18}.items(): pre.column_dimensions[c].width=w
    cur=3; blocks=[]
    def block(title,actname,personal,logistica):
        nonlocal cur
        _set(pre,f"A{cur}",title,fW,NAVY)
        for c in range(1,7): pre.cell(row=cur,column=c).fill=fill(NAVY); pre.cell(row=cur,column=c).border=BORD
        _set(pre,f"F{cur}",f'=IF({TOG(actname)}=1,"ACTIVA","INACTIVA")',fW,NAVY,al=CEN); cur+=1
        for i,h in enumerate(["Item","Dedicacion","Cantidad","Tiempo (meses)","Costo mensual ($)","Valor parcial ($)"]):
            _set(pre,f"{get_column_letter(i+1)}{cur}",h,fH,BLUE,al=CEN)
        cur+=1; first=cur
        for p in personal:
            if not p.get("cargo"): continue
            _set(pre,f"A{cur}",p["cargo"]); _set(pre,f"B{cur}",_num(p["dedicacion"]),fN,INPUT,NUM,CEN)
            _set(pre,f"C{cur}",_num(p["cantidad"]),fN,INPUT,al=CEN); _set(pre,f"D{cur}",f"={TMC(actname)}",fN,None,NUM,CEN)
            _set(pre,f"E{cur}",f"={CC(p['cargo'])}",fN,None,MONEY,RIG)
            _set(pre,f"F{cur}",f"=IFERROR(B{cur}*C{cur}*D{cur}*E{cur},0)",fN,None,MONEY,RIG); cur+=1
        for lg in logistica:
            if lg.get("concepto") in (None,""): continue
            val=lg["valor"]
            _set(pre,f"A{cur}",lg["concepto"]); _set(pre,f"C{cur}",_num(lg["cantidad"]),fN,INPUT,NUM,CEN)
            if isinstance(val,str) and val in LOG:
                _set(pre,f"E{cur}",f"={CL(val)}",fN,None,MONEY,RIG)   # referencia al catalogo
            else:
                _set(pre,f"E{cur}",_num(val),fN,INPUT,MONEY,RIG)      # monto global
            _set(pre,f"F{cur}",f"=IFERROR(C{cur}*E{cur},0)",fN,None,MONEY,RIG); cur+=1
        last=cur-1
        _set(pre,f"A{cur}",f"Subtotal — {title[:30]}",fB,GREY)
        for c in range(2,6): pre.cell(row=cur,column=c).fill=fill(GREY); pre.cell(row=cur,column=c).border=BORD
        _set(pre,f"F{cur}",f"=SUM(F{first}:F{last})",fB,GREY,MONEY,RIG); sub=cur; cur+=1
        _set(pre,f"A{cur}","Utilidad esperada",fN,GREY)
        for c in range(2,5): pre.cell(row=cur,column=c).fill=fill(GREY); pre.cell(row=cur,column=c).border=BORD
        _set(pre,f"E{cur}",f"={UTIL}",fN,GREY,PCT,CEN); _set(pre,f"F{cur}",f"=F{sub}*{UTIL}",fN,GREY,MONEY,RIG); util=cur; cur+=1
        _set(pre,f"A{cur}",f"Subtotal con utilidad — {title[:24]}",fB,GREEN)
        for c in range(2,6): pre.cell(row=cur,column=c).fill=fill(GREEN); pre.cell(row=cur,column=c).border=BORD
        _set(pre,f"F{cur}",f"=F{sub}+F{util}",fB,GREEN,MONEY,RIG); subu=cur; cur+=2
        blocks.append((title,sub,subu,TOG(actname)))
    for act in cfg["actividades"]:
        block(act["nombre"]+" — Fase de campo",act["nombre"],act["personal"],act["logistica"])
    # entregables
    _set(pre,f"A{cur}","ENTREGABLES — Fase de desarrollo",fW,NAVY)
    for c in range(1,7): pre.cell(row=cur,column=c).fill=fill(NAVY); pre.cell(row=cur,column=c).border=BORD
    _set(pre,f"F{cur}","ACTIVA",fW,NAVY,al=CEN); cur+=1
    for i,h in enumerate(["Entregable","Dedicacion","Cantidad","Tiempo (meses)","Costo mensual ($)","Valor parcial ($)"]):
        _set(pre,f"{get_column_letter(i+1)}{cur}",h,fH,BLUE,al=CEN)
    cur+=1; ef=cur
    for it in cfg["entregables"]["items"]:
        if not it.get("cargo"): continue
        _set(pre,f"A{cur}",it["nombre"]); _set(pre,f"B{cur}",_num(it["dedicacion"]),fN,INPUT,NUM,CEN)
        _set(pre,f"C{cur}",_num(it["cantidad"]),fN,INPUT,NUM,CEN); _set(pre,f"D{cur}",_num(it["tiempo_meses"]),fN,INPUT,NUM,CEN)
        _set(pre,f"E{cur}",f"={CC(it['cargo'])}",fN,None,MONEY,RIG)
        _set(pre,f"F{cur}",f"=IFERROR(B{cur}*C{cur}*D{cur}*E{cur},0)",fN,None,MONEY,RIG); cur+=1
    el=cur-1
    _set(pre,f"A{cur}","Subtotal — Entregables",fB,GREY)
    for c in range(2,6): pre.cell(row=cur,column=c).fill=fill(GREY); pre.cell(row=cur,column=c).border=BORD
    _set(pre,f"F{cur}",f"=SUM(F{ef}:F{el})",fB,GREY,MONEY,RIG); esub=cur; cur+=1
    _set(pre,f"A{cur}","Utilidad esperada",fN,GREY)
    for c in range(2,5): pre.cell(row=cur,column=c).fill=fill(GREY); pre.cell(row=cur,column=c).border=BORD
    _set(pre,f"E{cur}",f"={UTIL}",fN,GREY,PCT,CEN); _set(pre,f"F{cur}",f"=F{esub}*{UTIL}",fN,GREY,MONEY,RIG); eutil=cur; cur+=1
    _set(pre,f"A{cur}","Subtotal con utilidad — Entregables",fB,GREEN)
    for c in range(2,6): pre.cell(row=cur,column=c).fill=fill(GREEN); pre.cell(row=cur,column=c).border=BORD
    _set(pre,f"F{cur}",f"=F{esub}+F{eutil}",fB,GREEN,MONEY,RIG); esubu=cur; cur+=2
    blocks.append(("Entregables",esub,esubu,"1" if cfg["entregables"].get("activa",True) else "0"))
    # totales
    _set(pre,f"A{cur}","TOTALES DEL PROYECTO",Font(size=12,bold=True,color="FFFFFF"),NAVY)
    for c in range(1,7): pre.cell(row=cur,column=c).fill=fill(NAVY); pre.cell(row=cur,column=c).border=BORD
    cur+=1
    cd="+".join([f"F{s}*({t})" for _,s,su,t in blocks]); cu="+".join([f"F{su}*({t})" for _,s,su,t in blocks])
    _set(pre,f"A{cur}","Total costo directo (sin utilidad)",fB,LIGHT)
    for c in range(2,6): pre.cell(row=cur,column=c).fill=fill(LIGHT); pre.cell(row=cur,column=c).border=BORD
    _set(pre,f"F{cur}",f"={cd}",fB,LIGHT,MONEY,RIG); cur+=1
    _set(pre,f"A{cur}","Total + utilidad (antes de IVA)",fB,LIGHT)
    for c in range(2,6): pre.cell(row=cur,column=c).fill=fill(LIGHT); pre.cell(row=cur,column=c).border=BORD
    _set(pre,f"F{cur}",f"={cu}",fB,LIGHT,MONEY,RIG); CU=cur; cur+=1
    _set(pre,f"A{cur}","IVA",fN,LIGHT)
    for c in range(2,5): pre.cell(row=cur,column=c).fill=fill(LIGHT); pre.cell(row=cur,column=c).border=BORD
    _set(pre,f"E{cur}",f"={IVA}",fN,LIGHT,PCT,CEN); _set(pre,f"F{cur}",f"=F{CU}*{IVA}",fN,LIGHT,MONEY,RIG); IVAR=cur; cur+=1
    _set(pre,f"A{cur}","TOTAL GENERAL (con IVA)",Font(size=12,bold=True,color="FFFFFF"),NAVY)
    for c in range(2,6): pre.cell(row=cur,column=c).fill=fill(NAVY); pre.cell(row=cur,column=c).border=BORD
    _set(pre,f"F{cur}",f"=F{CU}+F{IVAR}",Font(size=12,bold=True,color="FFFFFF"),NAVY,MONEY,RIG)

    # ---- Tabla general ----
    tg=wb.create_sheet("Tabla general"); tg.sheet_view.showGridLines=False
    _set(tg,"A1","PROPUESTA ECONOMICA",Font(size=14,bold=True,color=NAVY),bd=False)
    _set(tg,"A2","=Parametros!B4",fB,bd=False)
    _set(tg,"A5","Esquema de desembolsos",fB,LIGHT); tg.merge_cells("A5:E5")
    for i,h in enumerate(["Desembolso","%","Valor antes de IVA ($)","IVA ($)","Total ($)"]): _set(tg,f"{get_column_letter(i+1)}6",h,fH,BLUE,al=CEN)
    dr=7
    for k in range(len(cfg["desembolsos"])):
        prow=DES_START+k
        _set(tg,f"A{dr}",f"=Parametros!A{prow}",fN,al=CEN); _set(tg,f"B{dr}",f"=Parametros!B{prow}",fN,None,PCT,CEN)
        _set(tg,f"C{dr}",f"=Presupuesto!$F${CU}*B{dr}",fN,None,MONEY,RIG)
        _set(tg,f"D{dr}",f"=C{dr}*{IVA}",fN,None,MONEY,RIG); _set(tg,f"E{dr}",f"=C{dr}+D{dr}",fN,None,MONEY,RIG); dr+=1
    _set(tg,f"A{dr}","Total",fB,GREY); _set(tg,f"B{dr}",f"=SUM(B7:B{dr-1})",fB,GREY,PCT,CEN)
    _set(tg,f"C{dr}",f"=SUM(C7:C{dr-1})",fB,GREY,MONEY,RIG); _set(tg,f"D{dr}",f"=SUM(D7:D{dr-1})",fB,GREY,MONEY,RIG); _set(tg,f"E{dr}",f"=SUM(E7:E{dr-1})",fB,GREY,MONEY,RIG)
    ir=dr+2
    _set(tg,f"A{ir}","Desglose por actividad",fB,LIGHT); tg.merge_cells(f"A{ir}:E{ir}"); ir+=1
    for i,h in enumerate(["Codigo","Item","Valor ($)","Valor + utilidad ($)","Valor por titulo ($)"]): _set(tg,f"{get_column_letter(i+1)}{ir}",h,fH,BLUE,al=CEN)
    ir+=1; istart=ir
    for idx,(label,sub,subu,t) in enumerate(blocks,1):
        _set(tg,f"A{ir}",f"A{idx}",fN,al=CEN); _set(tg,f"B{ir}",label)
        _set(tg,f"C{ir}",f"=Presupuesto!$F${sub}*({t})",fN,None,MONEY,RIG)
        _set(tg,f"D{ir}",f"=Presupuesto!$F${subu}*({t})",fN,None,MONEY,RIG)
        _set(tg,f"E{ir}",f"=IFERROR(D{ir}/{NTIT},0)",fN,None,MONEY,RIG); ir+=1
    iend=ir-1
    _set(tg,f"A{ir}","TOTAL",fB,GREY); _set(tg,f"B{ir}","",fB,GREY)
    _set(tg,f"C{ir}",f"=SUM(C{istart}:C{iend})",fB,GREY,MONEY,RIG); _set(tg,f"D{ir}",f"=SUM(D{istart}:D{iend})",fB,GREY,MONEY,RIG); _set(tg,f"E{ir}",f"=IFERROR(D{ir}/{NTIT},0)",fB,GREY,MONEY,RIG); td=ir; ir+=1
    _set(tg,f"A{ir}","IVA (19%)",fN,GREY); _set(tg,f"B{ir}","",fN,GREY)
    _set(tg,f"C{ir}",f"=C{td}*{IVA}",fN,GREY,MONEY,RIG); _set(tg,f"D{ir}",f"=D{td}*{IVA}",fN,GREY,MONEY,RIG); _set(tg,f"E{ir}",f"=IFERROR(D{ir}/{NTIT},0)",fN,GREY,MONEY,RIG); ivr=ir; ir+=1
    _set(tg,f"A{ir}","TOTAL GENERAL",fW,NAVY); _set(tg,f"B{ir}","",fW,NAVY)
    _set(tg,f"C{ir}",f"=C{td}+C{ivr}",fW,NAVY,MONEY,RIG); _set(tg,f"D{ir}",f"=D{td}+D{ivr}",fW,NAVY,MONEY,RIG); _set(tg,f"E{ir}",f"=IFERROR(D{ir}/{NTIT},0)",fW,NAVY,MONEY,RIG)
    tg.column_dimensions["A"].width=12; tg.column_dimensions["B"].width=42
    for col in "CDE": tg.column_dimensions[col].width=18
    wb.save(path)
    return path

# ============================ WORD (propuesta cliente) ============================
def _money(v): return "$ {:,.0f}".format(round(v))

def build_docx(cfg, budget, path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    NAVYC=RGBColor(0x1F,0x38,0x64); BLUEC=RGBColor(0x2E,0x75,0xB6)
    doc=Document()
    st=doc.styles["Normal"]; st.font.name="Calibri"; st.font.size=Pt(11)
    pj=cfg["proyecto"]
    def shade(cell,hexc):
        tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),hexc); tcPr.append(sh)
    # titulo
    h=doc.add_heading("",level=0)
    run=h.add_run("PROPUESTA TECNICA Y ECONOMICA"); run.font.color.rgb=NAVYC; run.font.size=Pt(20); run.bold=True
    p=doc.add_paragraph(); r=p.add_run(pj["objeto"]); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=BLUEC
    meta=doc.add_paragraph()
    meta.add_run(f"Codigo: {pj['codigo']}\n").bold=True
    if pj.get("cliente"): meta.add_run(f"Cliente: {pj['cliente']}\n")
    meta.add_run(f"Fecha: {pj['fecha']}\n")
    meta.add_run(f"Numero de titulos mineros: {pj['n_titulos']}")

    doc.add_heading("1. Propuesta economica",level=1)
    doc.add_paragraph("A continuacion se presenta el desglose economico de la oferta por actividad, "
                      "incluyendo utilidad e IVA. Los valores estan expresados en pesos colombianos (COP).")
    # tabla items
    activos=[b for b in budget["bloques"] if b["activa"]]
    t=doc.add_table(rows=1,cols=4); t.style="Light Grid Accent 1"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    hdr=t.rows[0].cells
    for i,txt in enumerate(["Codigo","Item","Valor + utilidad ($)","Valor por titulo ($)"]):
        hdr[i].text=txt
        for par_ in hdr[i].paragraphs:
            for rr in par_.runs: rr.bold=True; rr.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        shade(hdr[i],"1F3864")
    for idx,b in enumerate(activos,1):
        cells=t.add_row().cells
        cells[0].text=f"A{idx}"; cells[1].text=b["nombre"]
        cells[2].text=_money(b["subtotal_u"]); cells[3].text=_money(b["subtotal_u"]/budget["n_titulos"])
    # totales
    for lab,val,pt in [("TOTAL (antes de IVA)",budget["cu"],budget["cu"]/budget["n_titulos"]),
                       ("IVA (19%)",budget["iva"],budget["iva"]/budget["n_titulos"]),
                       ("TOTAL GENERAL",budget["total"],budget["total"]/budget["n_titulos"])]:
        cells=t.add_row().cells; cells[0].text=""; cells[1].text=lab
        cells[2].text=_money(val); cells[3].text=_money(pt)
        for c in cells:
            for par_ in c.paragraphs:
                for rr in par_.runs: rr.bold=True
            if lab=="TOTAL GENERAL":
                shade(c,"D9E2F3")

    doc.add_heading("2. Forma de pago (desembolsos)",level=1)
    t2=doc.add_table(rows=1,cols=5); t2.style="Light Grid Accent 1"; t2.alignment=WD_TABLE_ALIGNMENT.CENTER
    h2=t2.rows[0].cells
    for i,txt in enumerate(["#","%","Valor antes de IVA ($)","IVA ($)","Total ($)"]):
        h2[i].text=txt
        for par_ in h2[i].paragraphs:
            for rr in par_.runs: rr.bold=True; rr.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        shade(h2[i],"1F3864")
    for d in budget["desembolsos"]:
        cells=t2.add_row().cells
        cells[0].text=str(d["n"]); cells[1].text=f"{d['pct']*100:.0f}%"
        cells[2].text=_money(d["antes_iva"]); cells[3].text=_money(d["iva"]); cells[4].text=_money(d["total"])
    cells=t2.add_row().cells; cells[0].text=""; cells[1].text="100%"
    cells[2].text=_money(budget["cu"]); cells[3].text=_money(budget["iva"]); cells[4].text=_money(budget["total"])
    for c in cells:
        for par_ in c.paragraphs:
            for rr in par_.runs: rr.bold=True
    doc.add_paragraph()
    obs=doc.add_paragraph()
    obs.add_run("Observaciones de los desembolsos: ").bold=True
    obs.add_run("; ".join([f"({d['n']}) {d['obs']}" for d in budget["desembolsos"]])+".")
    nota=doc.add_paragraph(); nr=nota.add_run("Valores en pesos colombianos (COP). Utilidad e IVA incluidos segun se indica. "
        "Propuesta sujeta a las condiciones tecnicas del anexo.")
    nr.italic=True; nr.font.size=Pt(9); nr.font.color.rgb=RGBColor(0x59,0x59,0x59)
    doc.save(path)
    return path

# ============================ CLI ============================
if __name__ == "__main__":
    import sys, os
    cfg_path = sys.argv[1] if len(sys.argv) > 1 else None
    outdir = sys.argv[2] if len(sys.argv) > 2 else "."
    if cfg_path and os.path.exists(cfg_path):
        with open(cfg_path, encoding="utf-8") as fh: cfg = json.load(fh)
    else:
        cfg = copy.deepcopy(DEFAULT_CONFIG)
        print("(usando configuracion por defecto)")
    budget = compute_budget(cfg)
    cod = cfg["proyecto"]["codigo"]
    xlsx = os.path.join(outdir, f"{cod}_Presupuesto.xlsx")
    docx = os.path.join(outdir, f"{cod}_Propuesta.docx")
    build_excel(cfg, xlsx)
    build_docx(cfg, budget, docx)
    print("Generados:")
    print("  ", xlsx)
    print("  ", docx)
    print(f"Total antes de IVA: {_money(budget['cu'])}  |  IVA: {_money(budget['iva'])}  |  TOTAL: {_money(budget['total'])}")
